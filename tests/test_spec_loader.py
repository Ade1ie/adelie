"""tests/test_spec_loader.py — Unit tests for the Spec Loader module."""
from __future__ import annotations

import json
import tempfile
from pathlib import Path

import pytest

import adelie.config as cfg


@pytest.fixture
def tmp_workspace(tmp_path, monkeypatch):
    """Provide a temporary workspace for each test."""
    monkeypatch.setattr(cfg, "WORKSPACE_PATH", tmp_path)
    import adelie.kb.retriever as r
    monkeypatch.setattr(r, "WORKSPACE_PATH", tmp_path)
    monkeypatch.setattr(r, "INDEX_FILE", tmp_path / "index.json")
    r.ensure_workspace()
    return tmp_path


class TestConvertToMarkdown:
    def test_md_file_read_as_is(self, tmp_path):
        from adelie.spec_loader import convert_to_markdown

        md_file = tmp_path / "test.md"
        md_file.write_text("# Hello\n\nWorld", encoding="utf-8")

        result = convert_to_markdown(md_file)
        assert "# Hello" in result
        assert "World" in result

    def test_unsupported_extension_raises(self, tmp_path):
        from adelie.spec_loader import convert_to_markdown

        txt_file = tmp_path / "test.txt"
        txt_file.write_text("hello", encoding="utf-8")

        with pytest.raises(ValueError, match="Unsupported"):
            convert_to_markdown(txt_file)

    def test_nonexistent_file_raises(self, tmp_path):
        from adelie.spec_loader import convert_to_markdown

        with pytest.raises(FileNotFoundError):
            convert_to_markdown(tmp_path / "nope.md")

    def test_pdf_conversion(self, tmp_path):
        """Test PDF to Markdown conversion with a minimal PDF."""
        from adelie.spec_loader import convert_to_markdown

        try:
            from PyPDF2 import PdfWriter
        except ImportError:
            pytest.skip("PyPDF2 not installed")

        # Create a minimal PDF with text
        pdf_path = tmp_path / "test.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=612, height=792)
        with open(pdf_path, "wb") as f:
            writer.write(f)

        result = convert_to_markdown(pdf_path)
        assert "test" in result.lower()  # Contains filename stem
        assert "Converted from" in result

    def test_docx_conversion(self, tmp_path):
        """Test DOCX to Markdown conversion with a real DOCX."""
        from adelie.spec_loader import convert_to_markdown

        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        # Create a DOCX with content
        doc = Document()
        doc.add_heading("Test Heading", level=1)
        doc.add_paragraph("This is a test paragraph.")
        doc.add_paragraph("Another paragraph.")
        docx_path = tmp_path / "test.docx"
        doc.save(str(docx_path))

        result = convert_to_markdown(docx_path)
        assert "Test Heading" in result
        assert "test paragraph" in result

    def test_docx_with_table(self, tmp_path):
        """Test DOCX table conversion to Markdown table."""
        from adelie.spec_loader import convert_to_markdown

        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        doc = Document()
        doc.add_paragraph("Before table")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Header1"
        table.cell(0, 1).text = "Header2"
        table.cell(1, 0).text = "Cell1"
        table.cell(1, 1).text = "Cell2"
        docx_path = tmp_path / "table_test.docx"
        doc.save(str(docx_path))

        result = convert_to_markdown(docx_path)
        assert "Header1" in result
        assert "Cell1" in result
        assert "|" in result  # MD table syntax


class TestLoadSpec:
    def test_load_md_spec(self, tmp_workspace, tmp_path):
        from adelie.spec_loader import load_spec

        md_file = tmp_path / "my_spec.md"
        md_file.write_text("# My API Spec\n\nGET /users", encoding="utf-8")

        result = load_spec(md_file, tmp_workspace, category="logic")
        assert result.exists()
        assert result.name == "spec_my_spec.md"
        assert "My API Spec" in result.read_text(encoding="utf-8")

    def test_load_spec_registers_in_index(self, tmp_workspace, tmp_path):
        from adelie.spec_loader import load_spec
        import adelie.kb.retriever as r

        md_file = tmp_path / "api.md"
        md_file.write_text("# API", encoding="utf-8")

        load_spec(md_file, tmp_workspace, category="dependencies")
        index = r.get_index()

        assert "dependencies/spec_api.md" in index
        assert "spec" in index["dependencies/spec_api.md"]["tags"]

    def test_load_spec_custom_category(self, tmp_workspace, tmp_path):
        from adelie.spec_loader import load_spec

        md_file = tmp_path / "errors.md"
        md_file.write_text("# Known Errors", encoding="utf-8")

        result = load_spec(md_file, tmp_workspace, category="errors")
        assert "errors" in str(result)

    def test_load_docx_spec(self, tmp_workspace, tmp_path):
        from adelie.spec_loader import load_spec

        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        doc = Document()
        doc.add_heading("Architecture", level=1)
        doc.add_paragraph("Microservices pattern")
        docx_path = tmp_path / "arch.docx"
        doc.save(str(docx_path))

        result = load_spec(docx_path, tmp_workspace)
        assert result.exists()
        content = result.read_text(encoding="utf-8")
        assert "Architecture" in content
        assert "Microservices" in content


class TestListSpecs:
    def test_list_empty(self, tmp_workspace):
        from adelie.spec_loader import list_specs

        specs = list_specs(tmp_workspace)
        assert specs == []

    def test_list_after_load(self, tmp_workspace, tmp_path):
        from adelie.spec_loader import load_spec, list_specs

        md_file = tmp_path / "spec_a.md"
        md_file.write_text("# Spec A", encoding="utf-8")
        load_spec(md_file, tmp_workspace)

        specs = list_specs(tmp_workspace)
        assert len(specs) == 1
        assert specs[0]["name"] == "spec_spec_a"
        assert specs[0]["category"] == "logic"

    def test_list_shows_chunk_count(self, tmp_workspace, tmp_path):
        from adelie.spec_loader import load_spec, list_specs

        # Create a large spec that will be chunked
        md_file = tmp_path / "big.md"
        content = "# Big Spec\n\n"
        for i in range(15):
            content += f"## Section {i}\n\n" + "Details " * 300 + "\n\n"
        md_file.write_text(content, encoding="utf-8")
        load_spec(md_file, tmp_workspace)

        specs = list_specs(tmp_workspace)
        assert len(specs) == 1
        assert specs[0]["chunks"] > 0


class TestRemoveSpec:
    def test_remove_existing(self, tmp_workspace, tmp_path):
        from adelie.spec_loader import load_spec, remove_spec, list_specs

        md_file = tmp_path / "to_remove.md"
        md_file.write_text("# Remove me", encoding="utf-8")
        load_spec(md_file, tmp_workspace)

        assert len(list_specs(tmp_workspace)) == 1
        result = remove_spec(tmp_workspace, "spec_to_remove")
        assert result is True
        assert len(list_specs(tmp_workspace)) == 0

    def test_remove_nonexistent(self, tmp_workspace):
        from adelie.spec_loader import remove_spec

        result = remove_spec(tmp_workspace, "nonexistent")
        assert result is False

    def test_remove_chunked_spec_cleans_all(self, tmp_workspace, tmp_path):
        from adelie.spec_loader import load_spec, remove_spec, list_specs

        # Create large spec with chunks
        md_file = tmp_path / "chunked.md"
        content = "# Chunked\n\n"
        for i in range(12):
            content += f"## Part {i}\n\n" + "Info " * 300 + "\n\n"
        md_file.write_text(content, encoding="utf-8")
        load_spec(md_file, tmp_workspace)

        # Verify chunks exist
        logic_dir = tmp_workspace / "logic"
        chunks_before = list(logic_dir.glob("spec_chunked_chunk_*.md"))
        assert len(chunks_before) > 0

        # Remove
        result = remove_spec(tmp_workspace, "spec_chunked")
        assert result is True

        # All gone
        chunks_after = list(logic_dir.glob("spec_chunked_chunk_*.md"))
        assert len(chunks_after) == 0
        assert len(list_specs(tmp_workspace)) == 0
